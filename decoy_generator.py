from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

def add_hyperlink(paragraph, url, text):
    # This function inserts a clickable hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue underline (default link style)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_decoy_doc(filename="decoy.docx"):
    doc = Document()
    doc.add_heading('Confidential Report', level=0)

    p1 = doc.add_paragraph()
    p1.alignment = 1
    p1.add_run("This document contains sensitive project data.\n\n")

    doc.add_paragraph("Chart & Illustration:")
    doc.add_picture("images.png", width=Inches(3))

    doc.add_paragraph("\nClick below for complete access:")
    p_link = doc.add_paragraph()
    p_link.alignment = 1
    add_hyperlink(p_link, "http://localhost:5000/beacon?id=decoy1", "Click here to view full report")

    doc.add_paragraph("\nPrepared by: Cybersecurity Intern Team")
    doc.save(filename)
    print(f"✅ Decoy document '{filename}' created successfully.")

if __name__ == "__main__":
    create_decoy_doc()
